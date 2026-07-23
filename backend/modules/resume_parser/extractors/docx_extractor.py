from pathlib import Path

from docx import Document

from .base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extract text from DOCX files."""

    def extract(self, file_path: Path) -> str:
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        return "\n".join(paragraphs + table_text).strip()
